"""Paper processing service — URL fetching, PDF parsing, metadata extraction."""

import os
import re
import json
import uuid
import sqlite3
import asyncio
from datetime import datetime
from typing import Optional

DB_PATH = os.getenv("DB_PATH", "./db/xiaotiao.db")
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads", "papers")

os.makedirs(UPLOAD_DIR, exist_ok=True)


def _connect(db_path: Optional[str] = None):
    conn = sqlite3.connect(db_path or DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def extract_arxiv_id(url: str) -> Optional[str]:
    """Extract arXiv paper ID from URL."""
    patterns = [
        r'arxiv\.org/abs/(\d+\.\d+)',
        r'arxiv\.org/pdf/(\d+\.\d+)',
        r'arxiv\.org/abs/([\w\-]+/\d+)',
    ]
    for pattern in patterns:
        m = re.search(pattern, url)
        if m:
            return m.group(1)
    return None


async def fetch_arxiv_metadata(arxiv_id: str) -> dict:
    """Fetch paper metadata from arXiv API."""
    import httpx
    api_url = f"http://export.arxiv.org/api/query?id_list={arxiv_id}"

    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(api_url)
        resp.raise_for_status()

    import xml.etree.ElementTree as ET
    ns = {'atom': 'http://www.w3.org/2005/Atom'}
    root = ET.fromstring(resp.text)
    entry = root.find('atom:entry', ns)

    if entry is None:
        return {}

    title = entry.findtext('atom:title', '', ns).strip().replace('\n', ' ')
    abstract = entry.findtext('atom:summary', '', ns).strip()
    authors = [a.findtext('atom:name', '', ns) for a in entry.findall('atom:author', ns)]

    # Find PDF link
    pdf_url = None
    for link in entry.findall('atom:link', ns):
        if link.get('title') == 'pdf':
            pdf_url = link.get('href')

    return {
        'title': title,
        'abstract': abstract,
        'authors': json.dumps(authors),
        'arxiv_id': arxiv_id,
        'pdf_url': pdf_url or f"https://arxiv.org/pdf/{arxiv_id}.pdf",
    }


async def fetch_metadata_by_title(title: str) -> dict:
    """Try to find paper metadata (authors, DOI, abstract) via CrossRef, then OpenAlex."""
    import httpx
    meta = {}

    # ── CrossRef ──
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.get(
                "https://api.crossref.org/works",
                params={"query.bibliographic": title, "rows": 1, "sort": "relevance"},
                headers={"User-Agent": "xiaotiao/1.0 (mailto:xiaotiao@example.com)"},
            )
            if resp.status_code == 200:
                items = resp.json().get("message", {}).get("items", [])
                if items:
                    item = items[0]
                    cr_title = (item.get("title") or [""])[0]
                    # Only use if title similarity is reasonable
                    if cr_title and _title_similar(title, cr_title):
                        meta["doi"] = item.get("DOI", "")
                        # Authors
                        authors = item.get("author", [])
                        author_names = []
                        for a in authors:
                            given = a.get("given", "")
                            family = a.get("family", "")
                            author_names.append(f"{given} {family}".strip())
                        if author_names:
                            meta["authors"] = json.dumps(author_names)
                        # Abstract
                        abstract = item.get("abstract", "")
                        if abstract:
                            meta["abstract"] = re.sub(r"<[^>]+>", "", abstract).strip()
                        return meta
    except Exception as e:
        print(f"[paper_service] CrossRef metadata lookup failed: {e}")

    # ── OpenAlex fallback ──
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.get(
                "https://api.openalex.org/works",
                params={"search": title, "per_page": 1, "mailto": "xiaotiao@example.com"},
                headers={"User-Agent": "xiaotiao/1.0"},
            )
            if resp.status_code == 200:
                results = resp.json().get("results", [])
                if results:
                    work = results[0]
                    w_title = work.get("title", "")
                    if w_title and _title_similar(title, w_title):
                        meta["doi"] = (work.get("doi") or "").replace("https://doi.org/", "")
                        # Authors from authorships
                        authorships = work.get("authorships", [])
                        author_names = [a.get("author", {}).get("display_name", "") for a in authorships if a.get("author")]
                        if author_names:
                            meta["authors"] = json.dumps(author_names)
                        # Abstract from inverted index
                        inv = work.get("abstract_inverted_index")
                        if inv and isinstance(inv, dict):
                            positions = []
                            for word, pos_list in inv.items():
                                for pos in pos_list:
                                    positions.append((pos, word))
                            positions.sort()
                            meta["abstract"] = " ".join(w for _, w in positions)
                        return meta
    except Exception as e:
        print(f"[paper_service] OpenAlex metadata lookup failed: {e}")

    return meta


def _title_similar(t1: str, t2: str) -> bool:
    """Simple title similarity check — normalized lowercase prefix match."""
    a = re.sub(r"[^a-z0-9 ]", "", t1.lower()).strip()
    b = re.sub(r"[^a-z0-9 ]", "", t2.lower()).strip()
    if not a or not b:
        return False
    # Check if first 40 chars match, or one contains the other
    return a[:40] == b[:40] or a in b or b in a


async def process_paper_url(paper_id: str, url: str, db_path: Optional[str] = None):
    """Background task: fetch metadata for a paper URL and update the record."""
    conn = _connect(db_path)
    try:
        conn.execute(
            "UPDATE papers SET status='processing', updated_at=? WHERE id=?",
            (datetime.utcnow().isoformat(), paper_id)
        )
        conn.commit()

        arxiv_id = extract_arxiv_id(url)
        if arxiv_id:
            meta = await fetch_arxiv_metadata(arxiv_id)
            if meta:
                conn.execute(
                    """UPDATE papers SET title=?, abstract=?, authors=?, arxiv_id=?,
                       pdf_url=?, status='ready', updated_at=? WHERE id=?""",
                    (meta['title'], meta['abstract'], meta['authors'],
                     meta['arxiv_id'], meta['pdf_url'],
                     datetime.utcnow().isoformat(), paper_id)
                )
                conn.commit()
                return

        # Non-arXiv URL: try generic fetch
        import httpx
        async with httpx.AsyncClient(timeout=20, follow_redirects=True) as client:
            resp = await client.get(url)
            # Extract title from HTML
            title_match = re.search(r'<title>(.*?)</title>', resp.text, re.IGNORECASE | re.DOTALL)
            title = title_match.group(1).strip() if title_match else url

        # Auto-enrich with CrossRef/OpenAlex metadata
        enriched = await fetch_metadata_by_title(title)
        doi = enriched.get("doi", "")
        authors = enriched.get("authors", "")
        abstract = enriched.get("abstract", "")

        update_fields = ["title=?", "status='ready'", "updated_at=?"]
        update_values = [title, datetime.utcnow().isoformat()]
        if doi:
            update_fields.append("doi=?")
            update_values.append(doi)
        if authors:
            update_fields.append("authors=?")
            update_values.append(authors)
        if abstract:
            update_fields.append("abstract=?")
            update_values.append(abstract)

        update_values.append(paper_id)
        conn.execute(
            f"UPDATE papers SET {', '.join(update_fields)} WHERE id=?",
            tuple(update_values)
        )
        conn.commit()
    except Exception as e:
        conn.execute(
            "UPDATE papers SET status='failed', updated_at=? WHERE id=?",
            (datetime.utcnow().isoformat(), paper_id)
        )
        conn.commit()
        print(f"[paper_service] Error processing {paper_id}: {e}")
    finally:
        conn.close()


def process_paper_pdf(paper_id: str, pdf_path: str, db_path: Optional[str] = None):
    """Extract text from uploaded PDF and update paper record."""
    conn = _connect(db_path)
    try:
        conn.execute(
            "UPDATE papers SET status='processing', updated_at=? WHERE id=?",
            (datetime.utcnow().isoformat(), paper_id)
        )
        conn.commit()

        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        full_text = ""
        for page in doc:
            full_text += page.get_text() + "\n"
        doc.close()

        # Truncate to ~8000 words
        words = full_text.split()
        if len(words) > 8000:
            full_text = " ".join(words[:8000])

        # Extract title from first meaningful line
        lines = [l.strip() for l in full_text.split('\n') if l.strip()]
        title = lines[0][:200] if lines else os.path.basename(pdf_path)

        conn.execute(
            """UPDATE papers SET title=?, abstract=?, content_text=?, status='ready', updated_at=? WHERE id=?""",
            (title, full_text[:2000], full_text, datetime.utcnow().isoformat(), paper_id)
        )
        conn.commit()
    except Exception as e:
        conn.execute(
            "UPDATE papers SET status='failed', updated_at=? WHERE id=?",
            (datetime.utcnow().isoformat(), paper_id)
        )
        conn.commit()
        print(f"[paper_service] PDF processing error {paper_id}: {e}")
    finally:
        conn.close()


async def process_paper_docx(paper_id: str, docx_path: str, db_path: Optional[str] = None):
    """Extract structured text from Word document and optionally enhance with AI."""
    conn = _connect(db_path)
    try:
        conn.execute(
            "UPDATE papers SET status='processing', updated_at=? WHERE id=?",
            (datetime.utcnow().isoformat(), paper_id)
        )
        conn.commit()

        try:
            import docx
        except Exception as exc:
            raise RuntimeError("缺少 python-docx 依赖") from exc

        document = docx.Document(docx_path)

        # ── 1. Extract structured text with heading hierarchy ──
        structured_parts = []
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name or style_name == "title":
                structured_parts.append(f"# {text}")
            elif "heading 2" in style_name or "subtitle" in style_name:
                structured_parts.append(f"## {text}")
            elif "heading 3" in style_name:
                structured_parts.append(f"### {text}")
            elif "heading 4" in style_name:
                structured_parts.append(f"#### {text}")
            else:
                structured_parts.append(text)

        # ── 2. Extract tables ──
        for idx, table in enumerate(document.tables):
            table_lines = [f"\n[表格 {idx + 1}]"]
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                table_lines.append(" | ".join(cells))
            structured_parts.append("\n".join(table_lines))

        full_text = "\n".join(structured_parts)

        # Truncate to ~8000 words
        words = full_text.split()
        if len(words) > 8000:
            full_text = " ".join(words[:8000])

        # Basic title extraction from first heading or paragraph
        raw_title = os.path.basename(docx_path)
        for part in structured_parts[:5]:
            clean = part.lstrip("#").strip()
            if clean and len(clean) > 2:
                raw_title = clean[:200]
                break

        # ── 3. AI-enhanced structure recognition ──
        ai_title = raw_title
        ai_abstract = full_text[:2000]

        try:
            from services.llm import call_claude_json
            ai_prompt = (
                "你是一位学术论文分析专家。请分析以下 Word 文档提取的文本，"
                "识别并返回以下字段（JSON 格式）：\n"
                '{"title": "论文标题", "abstract": "论文摘要（200-500字概括）", '
                '"structured_content": "重新整理的结构化正文（Markdown 格式，保留标题层级）"}\n\n'
                "如果文本不是学术论文，则 title 取文档标题，abstract 取前几段的概括，"
                "structured_content 保持原文结构。\n\n"
                "重要：只返回 JSON，不要包含 markdown 代码块标记。"
            )
            result = await call_claude_json(ai_prompt, full_text[:6000], max_tokens=4000)
            if isinstance(result, dict):
                if result.get("title"):
                    ai_title = result["title"][:200]
                if result.get("abstract"):
                    ai_abstract = result["abstract"][:2000]
                if result.get("structured_content"):
                    full_text = result["structured_content"]
        except Exception as e:
            print(f"[paper_service] AI enhancement failed for {paper_id}, using raw text: {e}")

        conn.execute(
            """UPDATE papers SET title=?, abstract=?, content_text=?, status='ready', updated_at=? WHERE id=?""",
            (ai_title, ai_abstract, full_text, datetime.utcnow().isoformat(), paper_id)
        )
        conn.commit()
    except Exception as e:
        conn.execute(
            "UPDATE papers SET status='failed', updated_at=? WHERE id=?",
            (datetime.utcnow().isoformat(), paper_id)
        )
        conn.commit()
        print(f"[paper_service] DOCX processing error {paper_id}: {e}")
    finally:
        conn.close()


def get_paper_text(paper_id: str, db_path: Optional[str] = None) -> str:
    """Get the full text content of a paper for AI context."""
    conn = _connect(db_path)
    try:
        row = conn.execute(
            "SELECT abstract, pdf_path, docx_path, content_text FROM papers WHERE id=?",
            (paper_id,),
        ).fetchone()
        if not row:
            return ""

        content_text = row["content_text"] if "content_text" in row.keys() else None
        if content_text:
            return content_text

        # If we have a local PDF, extract fresh text
        if row['pdf_path'] and os.path.exists(row['pdf_path']):
            try:
                import fitz
                doc = fitz.open(row['pdf_path'])
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
                words = text.split()
                return " ".join(words[:8000])
            except Exception:
                pass

        docx_path = row["docx_path"] if "docx_path" in row.keys() else None
        if docx_path and os.path.exists(docx_path):
            try:
                import docx
                document = docx.Document(docx_path)
                parts = []
                for para in document.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = (para.style.name or "").lower()
                    if "heading 1" in style_name or style_name == "title":
                        parts.append(f"# {text}")
                    elif "heading 2" in style_name or "subtitle" in style_name:
                        parts.append(f"## {text}")
                    elif "heading 3" in style_name:
                        parts.append(f"### {text}")
                    else:
                        parts.append(text)
                for idx, table in enumerate(document.tables):
                    table_lines = [f"\n[表格 {idx + 1}]"]
                    for row_t in table.rows:
                        cells = [cell.text.strip().replace("\n", " ") for cell in row_t.cells]
                        table_lines.append(" | ".join(cells))
                    parts.append("\n".join(table_lines))
                text = "\n".join(parts)
                words = text.split()
                return " ".join(words[:8000])
            except Exception:
                pass

        # Fall back to stored abstract
        return row['abstract'] or ""
    finally:
        conn.close()
